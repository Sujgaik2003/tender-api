"""
HumanlyAI - Content Paraphraser & Humanizer
---------------------------------------------
Standalone humanization engine using:
1. Synonym replacement with context awareness
2. Sentence restructuring
3. Phrase-level paraphrasing
4. LLM-powered deep rewriting
5. Multi-pass refinement
"""

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from pydantic import BaseModel, Field
from typing import Optional, List, Tuple
import httpx
import re
import random
import pdfplumber
import docx
import io

from config import get_settings

settings = get_settings()
router = APIRouter()


# ============ Models ============

class HumanizeRequest(BaseModel):
    text: str = Field(..., description="Text to humanize/paraphrase")
    max_ai_percentage: float = Field(30.0, ge=0, le=100)
    max_attempts: int = Field(10, ge=1, le=20)
    style: str = Field("professional", description="professional, casual, formal, simple, academic")
    mode: str = Field("balanced", description="light, balanced, aggressive, creative")


class HumanizeResponse(BaseModel):
    success: bool
    original_text: str
    humanized_text: str
    original_ai_percentage: float
    final_ai_percentage: float
    attempts_used: int
    techniques_applied: List[str]
    message: str


# ============ SYNONYM DATABASE ============

SYNONYMS = {
    "use": ["employ", "apply", "utilize", "adopt", "make use of"],
    "utilize": ["use", "employ", "apply", "harness"],
    "leverage": ["use", "employ", "capitalize on", "take advantage of"],
    "implement": ["execute", "carry out", "put into practice", "deploy", "roll out"],
    "facilitate": ["enable", "help", "assist", "support", "make possible"],
    "enhance": ["improve", "boost", "strengthen", "elevate", "upgrade"],
    "optimize": ["improve", "refine", "fine-tune", "streamline"],
    "achieve": ["accomplish", "attain", "reach", "realize"],
    "ensure": ["guarantee", "make sure", "confirm", "secure"],
    "provide": ["offer", "supply", "deliver", "give", "furnish"],
    "enable": ["allow", "permit", "let", "empower", "make possible"],
    "demonstrate": ["show", "display", "illustrate", "exhibit", "prove"],
    "establish": ["set up", "create", "found", "build", "form"],
    "maintain": ["keep", "preserve", "sustain", "uphold", "retain"],
    "require": ["need", "demand", "call for", "necessitate"],
    "develop": ["create", "build", "design", "craft", "form"],
    "consider": ["think about", "examine", "look at", "review", "weigh"],
    "evaluate": ["assess", "analyze", "examine", "review", "judge"],
    "indicate": ["show", "suggest", "point to", "reveal", "signal"],
    "address": ["tackle", "handle", "deal with", "take on", "resolve"],
    "comprehensive": ["complete", "thorough", "full", "extensive", "all-inclusive"],
    "robust": ["strong", "solid", "sturdy", "resilient", "reliable"],
    "innovative": ["new", "novel", "creative", "original", "groundbreaking"],
    "significant": ["major", "important", "notable", "considerable", "substantial"],
    "essential": ["vital", "crucial", "key", "necessary", "fundamental"],
    "effective": ["successful", "productive", "efficient", "powerful"],
    "efficient": ["effective", "productive", "streamlined", "optimized"],
    "various": ["different", "diverse", "several", "numerous", "multiple"],
    "complex": ["complicated", "intricate", "sophisticated", "elaborate"],
    "critical": ["crucial", "vital", "key", "essential", "important"],
    "substantial": ["significant", "considerable", "major", "large", "sizable"],
    "pivotal": ["key", "crucial", "central", "vital", "essential"],
    "paramount": ["supreme", "top", "chief", "primary", "foremost"],
    "seamless": ["smooth", "effortless", "fluid", "uninterrupted"],
    "strategic": ["planned", "calculated", "deliberate", "tactical"],
    "methodology": ["method", "approach", "technique", "process", "system"],
    "framework": ["structure", "system", "model", "foundation"],
    "paradigm": ["model", "pattern", "example", "standard"],
    "endeavor": ["effort", "attempt", "undertaking", "venture"],
    "landscape": ["environment", "scene", "field", "arena", "sphere"],
    "synergy": ["cooperation", "teamwork", "collaboration", "partnership"],
    "implementation": ["execution", "rollout", "deployment", "application"],
    "utilization": ["use", "usage", "application", "employment"],
    "significantly": ["greatly", "considerably", "substantially", "noticeably"],
    "effectively": ["successfully", "well", "efficiently", "properly"],
    "subsequently": ["later", "afterward", "then", "next"],
    "predominantly": ["mainly", "mostly", "primarily", "largely"],
    "consequently": ["as a result", "therefore", "thus", "hence"],
    "additionally": ["also", "moreover", "plus", "besides", "as well"],
    "furthermore": ["also", "in addition", "plus", "what's more"],
    "however": ["but", "yet", "still", "though", "nevertheless"],
    "therefore": ["so", "thus", "hence", "as a result"],
    "moreover": ["also", "besides", "in addition", "plus"],
}

PHRASE_PARAPHRASES = {
    "it is important to note that": ["notably", "it's worth mentioning that", "keep in mind that", ""],
    "it should be noted that": ["note that", "importantly", "notably", ""],
    "in order to": ["to", "so as to", "for"],
    "due to the fact that": ["because", "since", "as"],
    "at this point in time": ["now", "currently", "at present"],
    "in the event that": ["if", "should", "in case"],
    "for the purpose of": ["to", "for", "in order to"],
    "with regard to": ["about", "regarding", "concerning", "on"],
    "in terms of": ["regarding", "concerning", "when it comes to", "for"],
    "a large number of": ["many", "numerous", "lots of", "plenty of"],
    "a significant amount of": ["much", "considerable", "substantial"],
    "on the other hand": ["however", "but", "conversely", "alternatively"],
    "as a result of": ["because of", "due to", "owing to", "from"],
    "in light of": ["given", "considering", "because of", "due to"],
    "with respect to": ["regarding", "about", "concerning", "for"],
    "in accordance with": ["following", "per", "as per", "according to"],
    "prior to": ["before", "ahead of", "preceding"],
    "subsequent to": ["after", "following", "post"],
    "in the process of": ["currently", "now", "busy with"],
    "take into consideration": ["consider", "think about", "account for"],
    "come to the conclusion": ["conclude", "decide", "determine"],
    "make a decision": ["decide", "choose", "determine"],
    "is able to": ["can", "is capable of"],
    "has the ability to": ["can", "is able to"],
    "there are many": ["many", "numerous", "several"],
    "it is possible that": ["possibly", "perhaps", "maybe"],
    "the fact that": ["that", "how"],
    "in spite of the fact that": ["although", "though", "despite"],
    "regardless of the fact that": ["although", "even though"],
}

SENTENCE_STARTER_VARIATIONS = {
    "This": ["The", "Such", "That", ""],
    "It is": ["", "There's", "We find that"],
    "There are": ["We have", "You'll find", ""],
    "The": ["This", "Our", "Their", ""],
    "We": ["Our team", "The company", ""],
    "They": ["The team", "These professionals", ""],
}


# ============ PARAPHRASING TECHNIQUES ============

def apply_synonym_replacement(text: str, intensity: float = 0.3) -> Tuple[str, int]:
    words = text.split()
    replacements = 0
    result = []

    for word in words:
        word_lower = word.lower().strip('.,!?;:')

        if word_lower in SYNONYMS and random.random() < intensity:
            synonyms = SYNONYMS[word_lower]
            replacement = random.choice(synonyms)

            if word[0].isupper():
                replacement = replacement.capitalize()

            trailing = ""
            for char in reversed(word):
                if char in '.,!?;:':
                    trailing = char + trailing
                else:
                    break

            result.append(replacement + trailing)
            replacements += 1
        else:
            result.append(word)

    return ' '.join(result), replacements


def apply_phrase_paraphrasing(text: str) -> Tuple[str, int]:
    result = text
    replacements = 0

    for phrase, alternatives in PHRASE_PARAPHRASES.items():
        pattern = re.compile(re.escape(phrase), re.IGNORECASE)
        matches = pattern.findall(result)

        if matches:
            replacement = random.choice(alternatives)
            result = pattern.sub(replacement, result, count=1)
            replacements += 1

    return result, replacements


def apply_sentence_restructuring(text: str) -> Tuple[str, int]:
    sentences = re.split(r'([.!?]+)', text)
    result = []
    changes = 0

    i = 0
    while i < len(sentences):
        sentence = sentences[i].strip()
        punct = sentences[i + 1] if i + 1 < len(sentences) else "."

        if not sentence:
            i += 2
            continue

        for starter, alternatives in SENTENCE_STARTER_VARIATIONS.items():
            if sentence.startswith(starter + " ") and random.random() > 0.6:
                alt = random.choice([a for a in alternatives if a])
                if alt:
                    sentence = alt + sentence[len(starter):]
                else:
                    rest = sentence[len(starter):].strip()
                    if rest and rest[0].islower():
                        sentence = rest[0].upper() + rest[1:]
                    else:
                        sentence = rest
                changes += 1
                break

        if " and " in sentence and random.random() > 0.7:
            if random.random() > 0.5:
                sentence = sentence.replace(" and ", " as well as ", 1)
            changes += 1

        if sentence.startswith("However") and random.random() > 0.5:
            sentence = "But" + sentence[7:]
            changes += 1
        elif sentence.startswith("Therefore") and random.random() > 0.5:
            sentence = "So" + sentence[9:]
            changes += 1
        elif sentence.startswith("Additionally") and random.random() > 0.5:
            sentence = "Also" + sentence[12:]
            changes += 1

        result.append(sentence + punct)
        i += 2

    return ' '.join(result), changes


def add_contractions(text: str) -> Tuple[str, int]:
    contractions = [
        (r"\bdo not\b", "don't"), (r"\bdoes not\b", "doesn't"),
        (r"\bcannot\b", "can't"), (r"\bwill not\b", "won't"),
        (r"\bshould not\b", "shouldn't"), (r"\bwould not\b", "wouldn't"),
        (r"\bcould not\b", "couldn't"), (r"\bis not\b", "isn't"),
        (r"\bare not\b", "aren't"), (r"\bwas not\b", "wasn't"),
        (r"\bwere not\b", "weren't"), (r"\bhas not\b", "hasn't"),
        (r"\bhave not\b", "haven't"), (r"\bhad not\b", "hadn't"),
        (r"\bit is\b", "it's"), (r"\bthat is\b", "that's"),
        (r"\bwhat is\b", "what's"), (r"\bwho is\b", "who's"),
        (r"\bthere is\b", "there's"), (r"\bhere is\b", "here's"),
        (r"\bthey are\b", "they're"), (r"\bwe are\b", "we're"),
        (r"\byou are\b", "you're"), (r"\bI am\b", "I'm"),
        (r"\bI have\b", "I've"), (r"\bI will\b", "I'll"),
        (r"\bI would\b", "I'd"), (r"\blet us\b", "let's"),
        (r"\bit will\b", "it'll"), (r"\bthat will\b", "that'll"),
    ]

    result = text
    count = 0

    for pattern, replacement in contractions:
        if random.random() > 0.3:
            before = result
            result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
            if before != result:
                count += 1

    return result, count


def remove_ai_markers(text: str) -> Tuple[str, int]:
    result = text
    count = 0

    ai_word_replacements = {
        "delve": "explore", "tapestry": "mix", "realm": "area",
        "landscape": "field", "journey": "process", "unlock": "discover",
        "empower": "enable", "seamless": "smooth", "robust": "strong",
        "holistic": "complete", "synergy": "cooperation", "paradigm": "approach",
        "innovative": "new", "cutting-edge": "modern", "state-of-the-art": "latest",
        "game-changer": "breakthrough", "groundbreaking": "major",
        "revolutionary": "significant", "unprecedented": "unique",
        "world-class": "excellent", "best-in-class": "top",
        "leverage": "use", "utilize": "use", "utilization": "use",
        "facilitate": "help", "comprehensive": "complete", "dynamic": "active",
        "evolving": "growing", "strategic": "planned", "pivotal": "key",
        "paramount": "vital", "endeavor": "effort", "noteworthy": "important",
        "fortify": "strengthen", "fostering": "encouraging",
        "bolster": "support", "underscore": "show",
        "multifaceted": "varied", "vibrant": "lively", "ongoing": "current",
    }

    for word, replacement in ai_word_replacements.items():
        pattern = re.compile(r'\b' + re.escape(word) + r'\b', re.IGNORECASE)
        if pattern.search(result):
            result = pattern.sub(replacement, result)
            count += 1

    ai_phrase_replacements = {
        "in essence": "", "at its core": "",
        "plays a crucial role": "is important", "plays a vital role": "matters",
        "plays an important role": "helps", "it's worth noting": "",
        "what's more": "also", "in today's world": "today",
        "in the modern era": "now", "continues to grow": "keeps growing",
        "continues to evolve": "keeps changing",
        "further reinforced": "strengthened",
        "continually growing": "growing", "continually adapting": "adapting",
    }

    for phrase, replacement in ai_phrase_replacements.items():
        pattern = re.compile(re.escape(phrase), re.IGNORECASE)
        if pattern.search(result):
            result = pattern.sub(replacement, result)
            count += 1

    result = re.sub(r'\s+', ' ', result).strip()
    result = re.sub(r'\s+([.,!?])', r'\1', result)

    return result, count


# ============ AI SCORE CALCULATION ============

def calculate_ai_score(text: str) -> Tuple[float, List[str]]:
    detected = []
    scores = {}
    text_lower = text.lower()
    words = text.split()
    word_count = len(words)

    if word_count < 5:
        return 0, ["text_too_short"]

    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip() and len(s.strip()) > 3]
    sentence_count = len(sentences)

    # 1. BURSTINESS
    burstiness_score = 0
    if sentence_count >= 2:
        sentence_lengths = [len(s.split()) for s in sentences]
        avg_len = sum(sentence_lengths) / len(sentence_lengths)
        if avg_len > 0:
            variance = sum((l - avg_len) ** 2 for l in sentence_lengths) / len(sentence_lengths)
            std_dev = variance ** 0.5
            cv = std_dev / avg_len
            if cv < 0.25:
                burstiness_score = 35
                detected.append("very_uniform_sentences")
            elif cv < 0.40:
                burstiness_score = 20
                detected.append("uniform_sentences")
            elif cv < 0.55:
                burstiness_score = 10
    scores['burstiness'] = burstiness_score

    # 2. LEXICAL DIVERSITY
    lexical_score = 0
    unique_words = set(w.lower().strip('.,!?;:"\'-') for w in words if len(w) > 2)
    if word_count > 0:
        ttr = len(unique_words) / word_count
        if ttr < 0.45:
            lexical_score = 25
            detected.append("low_vocabulary_diversity")
        elif ttr < 0.55:
            lexical_score = 15
        elif ttr < 0.65:
            lexical_score = 8
    scores['lexical'] = lexical_score

    # 3. FORMALITY
    formality_score = 0
    contraction_patterns = [
        r"\bdon't\b", r"\bcan't\b", r"\bwon't\b", r"\bisn't\b", r"\baren't\b",
        r"\bwasn't\b", r"\bweren't\b", r"\bhasn't\b", r"\bhaven't\b", r"\bit's\b",
        r"\bthat's\b", r"\bwhat's\b", r"\bthey're\b", r"\bwe're\b", r"\bI'm\b",
    ]
    has_contractions = any(re.search(p, text) for p in contraction_patterns)
    if word_count > 30 and not has_contractions:
        formality_score += 15
        detected.append("no_contractions")

    formal_transitions = [
        (r"\bfurthermore\b", 8), (r"\bmoreover\b", 8), (r"\bnevertheless\b", 8),
        (r"\bnonetheless\b", 8), (r"\bconsequently\b", 7), (r"\bsubsequently\b", 7),
        (r"\badditionall?y\b", 5), (r"\bhowever\b", 3), (r"\btherefore\b", 4),
        (r"\bthus\b", 5), (r"\bhence\b", 6), (r"\bin particular\b", 4),
    ]
    for pattern, penalty in formal_transitions:
        if re.search(pattern, text_lower):
            formality_score += penalty
    scores['formality'] = min(formality_score, 40)

    # 4. SENTENCE STARTERS
    starter_score = 0
    if sentence_count >= 3:
        starters = [s.split()[0].lower() for s in sentences if s.split()]
        starter_counts = {}
        for s in starters:
            starter_counts[s] = starter_counts.get(s, 0) + 1
        for starter, count in starter_counts.items():
            ratio = count / len(starters)
            if ratio >= 0.5:
                starter_score += 20
            elif ratio >= 0.33 and count >= 2:
                starter_score += 10
    scores['starters'] = min(starter_score, 25)

    # 5. AI BUZZWORDS
    buzzword_score = 0
    ai_buzzwords = [
        (r"\bdelve\b", 15), (r"\btapestry\b", 12), (r"\blandscape\b", 5),
        (r"\bjourney\b", 4), (r"\bunlock\b", 5), (r"\bempower\b", 5),
        (r"\bseamless\b", 6), (r"\brobust\b", 5), (r"\binnovative\b", 4),
        (r"\bholistic\b", 8), (r"\bsynergy\b", 10), (r"\bparadigm\b", 10),
        (r"\bcutting-edge\b", 8), (r"\bstate-of-the-art\b", 8),
        (r"\bgroundbreaking\b", 6), (r"\bunprecedented\b", 5),
        (r"\bpivotal\b", 6), (r"\bparamount\b", 7), (r"\bmultifaceted\b", 8),
    ]
    for pattern, penalty in ai_buzzwords:
        count = len(re.findall(pattern, text_lower))
        if count > 0:
            buzzword_score += penalty * min(count, 2)
    scores['buzzwords'] = min(buzzword_score, 30)

    # 6. AI PHRASES
    phrase_score = 0
    ai_phrases = [
        (r"it is important to note", 15), (r"it should be noted", 12),
        (r"it is worth mentioning", 12), (r"this highlights the", 6),
        (r"this underscores", 8), (r"in today's world", 8),
        (r"in the modern era", 8), (r"in conclusion", 6),
        (r"plays a crucial role", 8), (r"plays a vital role", 8),
        (r"at its core", 6), (r"at the heart of", 6),
    ]
    for phrase, penalty in ai_phrases:
        if re.search(phrase, text_lower):
            phrase_score += penalty
    scores['phrases'] = min(phrase_score, 35)

    # FINAL SCORE
    total_raw = sum(scores.values())
    if word_count < 50:
        final_score = min(total_raw * 0.8, 100)
    elif word_count < 100:
        final_score = min(total_raw * 0.9, 100)
    else:
        final_score = min(total_raw, 100)

    return round(final_score, 2), detected


# ============ LLM PARAPHRASING ============

async def llm_paraphrase(text: str, style: str, mode: str, attempt: int) -> str:
    if not settings.llm_api_key:
        result = text
        result, _ = apply_synonym_replacement(result, 0.4)
        result, _ = apply_phrase_paraphrasing(result)
        result, _ = add_contractions(result)
        return result

    mode_instructions = {
        "light": "Make minimal changes. Only fix obvious AI patterns.",
        "balanced": "Rewrite moderately while staying close to original length.",
        "aggressive": "Completely rewrite while keeping it SHORT and simple.",
        "creative": "Rewrite naturally like a real person would speak.",
    }

    style_instructions = {
        "professional": "business-like but natural",
        "casual": "friendly, like texting a colleague",
        "formal": "proper but not robotic",
        "simple": "simple words, short sentences",
        "academic": "scholarly but readable",
    }

    original_word_count = len(text.split())

    prompt = f"""Rewrite this COMPLETE text to sound human-written, not AI-generated.

STRICT RULES (MUST FOLLOW):
1. Keep EXACTLY the same meaning - do not omit or summarize any information.
2. Maintain the same depth and detail as the original.
3. Use contractions where natural: don't, can't, it's, they're, we're.
4. Mix sentence lengths for a natural rhythm - some short, some longer.
5. Every single point in the original must be present in your rewrite.

WORDS TO AVOID (replace with simpler alternatives):
- leverage → use, utilize → use, facilitate → help
- comprehensive → full/complete, robust → strong, innovative → new
- furthermore/moreover → also/plus, additionally → also

PHRASES TO REMOVE COMPLETELY:
- "It is important to note that"
- "It should be noted"
- "In essence"
- "plays a crucial role"
- "at its core"

STYLE: {style_instructions.get(style, style_instructions['professional'])}
MODE: {mode_instructions.get(mode, mode_instructions['balanced'])}

TEXT TO REWRITE:
{text}

REWRITTEN DOCUMENT (matching original length of approximately {original_word_count} words):"""

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            safe_max_tokens = max(1000, int(original_word_count * 2.5) + 500)

            response = await client.post(
                f"{settings.llm_api_url.rstrip('/')}/chat/completions",
                headers={"Authorization": f"Bearer {settings.llm_api_key}"},
                json={
                    "model": settings.llm_model,
                    "messages": [{"role": "user", "content": prompt}],
                    "max_tokens": safe_max_tokens,
                    "temperature": 0.6 + (attempt * 0.08),
                }
            )

            if response.status_code == 200:
                result = response.json()
                paraphrased = result["choices"][0]["message"]["content"].strip()

                for prefix in ["Here's the paraphrased", "Paraphrased:", "Here is",
                               "Here's the rewritten", "Rewritten:", "REWRITTEN:",
                               "Rewritten Document:"]:
                    if paraphrased.lower().startswith(prefix.lower()):
                        paraphrased = paraphrased[len(prefix):].strip()
                        if paraphrased.startswith(":"):
                            paraphrased = paraphrased[1:].strip()

                if paraphrased.startswith('"') and paraphrased.endswith('"'):
                    paraphrased = paraphrased[1:-1]

                return paraphrased
    except Exception as e:
        print(f"[PARAPHRASE] LLM Error: {e}")

    # Fallback
    result = text
    result, _ = apply_synonym_replacement(result, 0.5)
    result, _ = apply_phrase_paraphrasing(result)
    result, _ = add_contractions(result)
    return result


# ============ MAIN HUMANIZE ENDPOINT ============

@router.post("/humanizer")
async def humanize_unified(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    max_ai_percentage: float = Form(30.0),
    max_attempts: int = Form(10),
    style: str = Form("professional"),
    mode: str = Form("balanced")
):
    """
    Unified endpoint for Humanizer.
    Accepts EITHER a file (PDF, DOCX, TXT) OR raw text.
    """
    content = ""

    # 1. Handle File
    if file:
        filename = file.filename.lower()
        try:
            file_bytes = await file.read()

            if filename.endswith(".pdf"):
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    content = "\n".join([page.extract_text() or "" for page in pdf.pages])
            elif filename.endswith(".docx") or filename.endswith(".doc"):
                if filename.endswith(".doc"):
                    raise HTTPException(status_code=400, detail="Legacy .doc files are not supported. Please save as .docx.")
                doc = docx.Document(io.BytesIO(file_bytes))
                content = "\n".join([p.text for p in doc.paragraphs])
            elif filename.endswith(".txt"):
                content = file_bytes.decode("utf-8", errors='ignore')
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported file: {filename}. Please use PDF, DOCX, or TXT.")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error parsing file: {str(e)}")

    # 2. Handle Text
    if text and text.strip():
        content = text

    if not content or not content.strip():
        raise HTTPException(status_code=400, detail="No content provided. Please upload a file or paste text.")

    if len(content) > 50000:
        content = content[:50000]

    original_text = content.strip()
    techniques_applied = []

    # Calculate original AI score
    original_ai_pct, _ = calculate_ai_score(original_text)
    print(f"[HUMANIZE] Original AI score: {original_ai_pct:.1f}%")

    if original_ai_pct <= max_ai_percentage:
        return {
            "transformed": original_text,
            "original_score": round(original_ai_pct, 2),
            "new_score": round(original_ai_pct, 2),
            "reduction": "0%"
        }

    # STEP 1: Rule-based transforms
    current_text = original_text
    current_text, n = remove_ai_markers(current_text)
    if n > 0: techniques_applied.append(f"ai_marker_removal:{n}")

    current_text, n = apply_phrase_paraphrasing(current_text)
    if n > 0: techniques_applied.append(f"phrase_paraphrase:{n}")

    intensity = {"light": 0.2, "balanced": 0.35, "aggressive": 0.5, "creative": 0.45}.get(mode, 0.35)
    current_text, n = apply_synonym_replacement(current_text, intensity)
    if n > 0: techniques_applied.append(f"synonym_replace:{n}")

    current_text, n = apply_sentence_restructuring(current_text)
    if n > 0: techniques_applied.append(f"restructure:{n}")

    current_text, n = add_contractions(current_text)
    if n > 0: techniques_applied.append(f"contractions:{n}")

    current_text = re.sub(r'\s+', ' ', current_text).strip()
    current_text = re.sub(r'\s+([.,!?])', r'\1', current_text)

    current_ai_pct, _ = calculate_ai_score(current_text)
    best_text = current_text
    best_ai_pct = current_ai_pct
    attempts_used = 0

    # STEP 2: LLM paraphrasing if needed
    if current_ai_pct > max_ai_percentage:
        techniques_applied.append("llm_paraphrase")

        for attempt in range(max_attempts):
            attempts_used = attempt + 1
            print(f"[HUMANIZE] LLM attempt {attempts_used}...")

            paraphrased = await llm_paraphrase(best_text, style, mode, attempt)
            paraphrased, _ = remove_ai_markers(paraphrased)
            paraphrased, _ = apply_phrase_paraphrasing(paraphrased)
            paraphrased, _ = add_contractions(paraphrased)
            paraphrased = re.sub(r'\s+', ' ', paraphrased).strip()

            new_ai_pct, _ = calculate_ai_score(paraphrased)
            print(f"[HUMANIZE] Attempt {attempts_used}: {new_ai_pct:.1f}%")

            if new_ai_pct < best_ai_pct:
                best_text = paraphrased
                best_ai_pct = new_ai_pct

            if best_ai_pct <= max_ai_percentage:
                break

    return {
        "transformed": best_text,
        "original_score": round(original_ai_pct, 2),
        "new_score": round(best_ai_pct, 2),
        "reduction": f"{original_ai_pct - best_ai_pct:.1f}%"
    }


# ============ TEXT EXTRACTION ENDPOINT ============

@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    """Extract text from uploaded file."""
    filename = file.filename.lower()
    file_bytes = await file.read()
    content = ""

    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                content = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(file_bytes))
            content = "\n".join([p.text for p in doc.paragraphs])
        elif filename.endswith(".txt"):
            content = file_bytes.decode("utf-8", errors='ignore')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error: {str(e)}")

    return {"text": content.strip()}


# ============ HEALTH CHECK ============

@router.get("/health")
async def health():
    return {
        "status": "healthy",
        "service": "humanly-ai-backend",
        "llm_available": bool(settings.llm_api_key),
    }
